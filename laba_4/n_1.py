from docx import Document as CreateDocument
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

doc: Document = CreateDocument()

doc.add_heading("Shrek Script.")
doc.add_heading("Written by Ted Elliot", level=2)
doc.add_heading("Based on book by William Steig", level=3)

p1 = doc.add_paragraph("Once upon a time there was a lovely princess. But she had an enchantment upon her of a fearful sort which could only be broken by love's first kiss. She was locked away in a castle guarded by a terrible fire-breathing dragon.\n"
    "Many brave knights had attempted to free her from this dreadful prison, but non prevailed. She waited in the dragon's keep in the highest room of the tallest tower for her true love and true love's first kiss. (laughs) Like that's ever gonna happen. What a load of -\n"
    "Allstar - by Smashmouth begins to play. Shrek goes about his day. While in a nearby town, the villagers get together to go after the ogre.\n"
)
p2 = doc.add_paragraph("There is a line of fairy tale creatures. The head of the guard sits at a table paying people for bringing the fairy tale creatures to him. There are cages all around. Some of the people in line are Peter Pan, who is carrying Tinkerbell in a cage, Gipetto who's carrying Pinocchio, and a farmer who is carrying the three little pigs.")

p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.first_line_indent = Mm(29)
p1.paragraph_format.left_indent = Mm(18.49)
p1.paragraph_format._line_spacing = Mm(0.04)

p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
p2.paragraph_format.page_break_before = True
p2.paragraph_format.right_indent = Mm(-0.004324)
p2.paragraph_format.left_indent = Mm(95)

img = doc.add_picture("img.jpeg", width=Mm(100), height=Mm(100))

doc.save("Amogus.gg.docx")